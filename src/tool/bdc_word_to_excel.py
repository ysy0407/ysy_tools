#!/usr/bin/env python3
# -*- coding:utf-8 -*-
# @Author: ysy
# @Time: 2022-03-18 14:57

import os
import re
import time
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font, colors
from src.util.path_util import base_path_join
from src.util.logger_util import logger

# 匹配表名的正则表达式
TABLE_NAME_PATTERN = '^([a-zA-Z0-9_]{3,})(\t| )*(（|\(){1}(.*)(\)|）){1}$'
# 开始获取表字段的表的序号，第0个为文档修订页的表格，故跳过
COLUMN_TABLE_START_INDEX = 1
# 结果excel的模板文件，第一个sheet是table_list，第二个sheet是表字段
RESULT_EXCEL_TEMPLATE_FILE_NAME = '数据处理接口数据说明文档-模板.xlsx'
RESULT_EXCEL_TEMPLATE_FILE_PATH = base_path_join('static', 'file', RESULT_EXCEL_TEMPLATE_FILE_NAME)


# 表信息
class TableInfo(object):

    def __init__(self, system_name, table_en_name, table_zh_name, column_table):
        self.system_name = system_name
        self.table_en_name = table_en_name
        self.table_zh_name = table_zh_name
        self.column_table = column_table

    def get_table_list_row(self):
        #       所属系统           表名                        描述                         所属业务功能模块  是/否生效 增/全量 表约束条件，增量抽取条件
        return [self.system_name, self.table_en_name.strip(), self.table_zh_name.strip(), '',            '是',    '全量', '',       '']

    def get_column_list_rows(self):
        column_list_rows = list()
        for i in range(len(self.column_table.rows)):
            # 第0行为word的表格里的表头，故跳过
            if i == 0:
                continue
            cells = self.column_table.rows[i].cells
            #                       所属表名字	        字段编号	字段英文名称	           字段中文名称	          字段类型	             长度	        主外键/索引标识	是否为空	字段转换	描述
            column_list_rows.append([self.table_en_name, i,     cells[0].text.strip(), cells[1].text.strip(), cells[2].text.strip(), cells[3].text, '',             '',     '',     cells[4].text])
        return column_list_rows

    def __str__(self) -> str:
        return 'table_en_name: ' + self.table_en_name \
               + ', table_zh_name: ' + self.table_zh_name \
               + ', len(column_table.rows): ' + str(len(self.column_table.rows) - 1)


# 处理类
class Handle(object):

    def __init__(self, system_name, word_file_path):
        self.system_name = system_name
        self.word_file_path = word_file_path
        self.excel_file_path = self.word_file_path.replace(
            os.path.split(self.word_file_path)[1],
            RESULT_EXCEL_TEMPLATE_FILE_NAME.replace('模板', self.system_name)
        )
        self.table_info_list = list()

    # 检查此段内容是否为表名，若为表名，返回(英文表名，中文表名)
    @staticmethod
    def check_is_table_name(paragraph):
        match_result = re.match(TABLE_NAME_PATTERN, paragraph)
        if match_result is None:
            return None
        else:
            return match_result.group(1), match_result.group(4)

    # 从Word文件中读取表信息list
    def read_table_info_list_from_word_file(self):
        document = Document(self.word_file_path)
        logger.info('system_name: %s, document.tables: %s', self.system_name, len(document.tables))
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text
            # 若段落为表名，则获取表名和描述
            check_result = self.check_is_table_name(paragraph_text)
            if check_result is not None:
                self.table_info_list.append(TableInfo(
                    self.system_name,
                    check_result[0],
                    check_result[1],
                    document.tables[len(self.table_info_list) - 1 + COLUMN_TABLE_START_INDEX]
                ))
                # print(len(self.table_info_list), self.table_info_list[-1])
        logger.info('len(table_info_list): %s', len(self.table_info_list))

    # 将表信息list写入到excel中
    def write_table_info_list_to_excel(self):
        workbook = load_workbook(RESULT_EXCEL_TEMPLATE_FILE_PATH)
        sheets = workbook.worksheets
        for i in range(len(self.table_info_list)):
            table_info = self.table_info_list[i]
            logger.info('handle index: %s, table info: %s', i + 1, table_info)
            # 向第0个sheet（TABLE_LIST）中增加当前表信息
            sheets[0].append(table_info.get_table_list_row())
            # 设置当前行表名的超链接和对应样式，跳转到对应的sheet
            table_en_name_cell = sheets[0].cell(row=i+2, column=2)
            table_en_name_cell.value = '=HYPERLINK("#{0}!A1","{1}")'.format(table_info.table_en_name, table_info.table_en_name)
            table_en_name_cell.font = Font(u='single', color=colors.BLUE)
            # 根据第1个sheet（TABLE_COLUMNS_TEMPLATE）复制出需要写入表字段的sheet
            table_columns_sheet = workbook.copy_worksheet(sheets[1])
            table_columns_sheet.title = table_info.table_en_name
            for column_list_row in table_info.get_column_list_rows():
                table_columns_sheet.append(column_list_row)
        workbook.remove(sheets[1])
        workbook.save(self.excel_file_path)
        pass

    def run(self):
        logger.info('handle starting')
        start_time = int(round(time.time() * 1000))
        self.read_table_info_list_from_word_file()
        self.write_table_info_list_to_excel()
        logger.info("handle over, use time: %sms", int(round(time.time() * 1000)) - start_time)
        return self.excel_file_path


def execute(params: dict):
    # 执行获取结果，返回结果文件路径
    return Handle(params["system_name"], params["word_file_path"]).run()


'''
    根据传入的系统名称和数据说明的docx文件，转为根据模板生成的excel文件
'''
if __name__ == '__main__':
    execute({
        'system_name': '电票',
        'word_file_path': r'C:\Users\什锦小沐\Desktop\城银清算\卸数平台\电票\票据交易系统-数据结构说明文档（导出版）.docx'
        # 'word_file_path': r'C:\Users\什锦小沐\Desktop\城银清算\卸数平台\python读取\票据交易系统-数据说明-模板.docx'

    })
    pass
